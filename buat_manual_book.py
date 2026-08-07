import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_manual_book():
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("MANUAL BOOK & PANDUAN PENGGUNAAN SISTEM\nRAJAWALI MOTOR SURABAYA")
    run_title.bold = True
    run_title.font.size = Pt(22)
    run_title.font.color.rgb = RGBColor(0xB0, 0x18, 0x1C)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Sistem Operasional, POS Kasir, Stok Multi-Barcode & Service Bengkel Motor\nVersi 2.0 (Laravel 13 - Production Grade)")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Heading 1 helper
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xB0, 0x18, 0x1C)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    def add_screenshot_box(title, desc):
        tbl = doc.add_table(rows=2, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False

        # Header cell
        cell_hdr = tbl.cell(0, 0)
        set_cell_background(cell_hdr, "F1F5F9")
        p_hdr = cell_hdr.paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_hdr = p_hdr.add_run(f"📸 SCREENSHOT: {title}")
        r_hdr.bold = True
        r_hdr.font.size = Pt(10)
        r_hdr.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

        # Body cell
        cell_body = tbl.cell(1, 0)
        set_cell_background(cell_body, "FAFAFA")
        p_body = cell_body.paragraphs[0]
        p_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_body = p_body.add_run(f"\n[ AREA TEMPAT FOTO SCREENSHOT / GAMBAR TAMPILAN ]\n\n{desc}\n")
        r_body.italic = True
        r_body.font.size = Pt(10)
        r_body.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Bab 1
    add_h1("1. PENDAHULUAN & RINGKASAN PERUBAHAN SISTEM")
    doc.add_paragraph(
        "Sistem Operasional Rajawali Motor Surabaya adalah platform manajemen bengkel motor terintegrasi "
        "yang mencakup Kasir Point of Sales (POS), Stok Multi-Barcode, Work Order Service Kendaraan, "
        "hingga Pengaturan Toko & Hak Akses Peran. Seluruh antarmuka dirancang modern, responsif untuk layar HP/Tablet/Desktop, "
        "serta dilengkapi dengan notifikasi interaktif SweetAlert2."
    )

    add_h2("Rincian Perubahan & Peningkatan Terbaru (Update v2.0):")
    bullet_points = [
        "Website Utama (Public Landing Page): Ditambahkan Smart Hide-on-Scroll Navbar interaktif, Floating Scroll-to-Top Button, dan Floating WhatsApp FAB.",
        "Kalkulator Estimasi Biaya: Mendukung kategori Sepeda Motor (Matic, Bebek, Maxi 250cc) & Mobil (City Car, SUV/MPV, Sedan) dengan perhitungan instant.",
        "Portal Login Split-Screen: Halaman login bertema White & Rajawali Red (#B0181C) lengkap dengan 1-Click Auto Fill Demo Role.",
        "Kasir POS Realtime: Terhubung ke database transaksi (Penjualan & PenjualanDetail), otomatis mengurangi stok barang pada StokMutasi, dan terintegrasi dengan Pop-up Cetak Struk Thermal 58mm/80mm.",
        "Global SweetAlert2 & Toast: Seluruh notifikasi simpan/update data menggunakan SweetAlert Toast di pojok kanan atas, serta modal konfirmasi hapus/nonaktifkan data.",
        "Manajemen Pengguna (User CRUD): Modul /admin/pengaturan/user untuk kelola akun, peran (Owner, Admin, Kasir, Gudang, Montir), dan reset password.",
        "40 Automated Tests: 100% Lulus pengujian otomatis (Unit & Feature Test)."
    ]
    for pt in bullet_points:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(pt)

    # Bab 2
    add_h1("2. DATA SEED AKUN DEMO & MATRIKS HAK AKSES PERAN")
    doc.add_paragraph(
        "Sistem memiliki 5 tingkatan peran (roles) untuk menjaga keamanan dan pembagian tugas operasional bengkel secara tepat:"
    )

    tbl_roles = doc.add_table(rows=6, cols=4)
    tbl_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Peran (Role)", "Username", "Password", "Hak Akses & Wewenang Utama"]
    
    for i, h in enumerate(headers):
        cell = tbl_roles.cell(0, i)
        set_cell_background(cell, "B0181C")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data_roles = [
        ("Owner", "owner", "password", "Akses Penuh 100% (Pengaturan Toko, User, Keuangan, Kasir, Stok, Laporan, Audit)"),
        ("Admin", "admin", "password", "Kelola Master Data Barang, Customer, Supplier, Sales, Stok, Penjualan, Pembelian"),
        ("Kasir", "kasir1", "password", "Aplikasi Kasir POS, Cetak Struk, Riwayat Penjualan, Data Customer"),
        ("Gudang", "gudang1", "password", "Master Barang, Barcode Multi, Stok Mutasi, Opname, Supplier, Pembelian"),
        ("Montir", "montir1", "password", "Work Order Service Kendaraan, Status Pengerjaan, Cetak Tanda Terima")
    ]

    for row_idx, data in enumerate(data_roles, start=1):
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = tbl_roles.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            p = cell.paragraphs[0]
            if col_idx in [1, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.font.name = 'Consolas'
                r.font.size = Pt(10)
            else:
                r = p.add_run(text)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Bab 3
    add_h1("3. ALUR KERJA UTAMA (OPERATIONAL WORKFLOWS)")

    add_h2("Alur 1: Autentikasi & Masuk Sistem (1-Click Auto Fill)")
    doc.add_paragraph("1. Buka URL `/admin/login` di browser.")
    doc.add_paragraph("2. Klik salah satu tombol '1-Click Auto Fill Demo' (Owner, Admin, Kasir, Gudang, atau Montir).")
    doc.add_paragraph("3. Klik tombol 'Masuk Ke System'. Sistem akan mengarahkan ke halaman utama sesuai peran masing-masing.")

    add_screenshot_box("Halaman Portal Login", "Menampilkan split-screen login dengan tombol 1-Click Auto Fill Demo Role dan wadah logo squircle emblem.")

    add_h2("Alur 2: Transaksi Penjualan POS Kasir & Cetak Struk")
    doc.add_paragraph("1. Buka menu Kasir (`/admin/kasir`).")
    doc.add_paragraph("2. Scan barcode barang atau ketik kode/nama barang pada kolom input utama (Shortcut: F2). Press Enter.")
    doc.add_paragraph("3. Pilih Customer (Umum atau Customer Terdaftar) dan Jenis Pembayaran (Tunai / Tempo).")
    doc.add_paragraph("4. Masukkan jumlah Uang Bayar (Shortcut: F9).")
    doc.add_paragraph("5. Tekan tombol Simpan Nota atau tekan `F12` / `Ctrl + Enter`.")
    doc.add_paragraph("6. Pop-up SweetAlert2 akan muncul mengonfirmasi keberhasilan transaksi. Klik 'Cetak Struk (58/80mm)' untuk mencetak ke printer thermal.")

    add_screenshot_box("Aplikasi Kasir POS & Pop-up Cetak Struk", "Menampilkan antarmuka kasir cepat, input barcode, keranjang barang, dan dialog cetak nota thermal.")

    add_h2("Alur 3: Manajemen Barang & Multi-Barcode Scanner")
    doc.add_paragraph("1. Buka menu Master Barang (`/admin/barang`).")
    doc.add_paragraph("2. Klik 'Tambah Barang Baru' untuk memasukkan kode, nama, group, satuan, HPP, dan harga eceran/grosir.")
    doc.add_paragraph("3. Klik ikon Barcode pada tabel untuk menambahkan barcode alternatif/vendor.")
    doc.add_paragraph("4. Setiap kali barang tersimpan, SweetAlert Toast akan muncul di pojok kanan atas.")

    add_screenshot_box("Halaman Master Barang & Kelola Barcode", "Menampilkan tabel master barang, badge status aktif, dan modal kelola multi-barcode.")

    add_h2("Alur 4: Manajemen User & Hak Akses Peran")
    doc.add_paragraph("1. Buka menu Pengaturan -> User (`/admin/pengaturan/user`) sebagai Owner.")
    doc.add_paragraph("2. Klik 'Tambah User Baru' untuk mendaftarkan akun baru lengkap dengan memilih Peran (Owner, Admin, Kasir, Gudang, Montir).")
    doc.add_paragraph("3. Klik ikon Ban/Check pada tabel untuk mengaktifkan/menonaktifkan akun.")

    add_screenshot_box("Halaman Manajemen User", "Menampilkan daftar pengguna sistem, badge peran berwarna khusus, dan modal konfirmasi SweetAlert2.")

    add_h2("Alur 5: Pengaturan Identitas Toko & Kebijakan Kasir")
    doc.add_paragraph("1. Buka menu Pengaturan -> Toko (`/admin/pengaturan/toko`).")
    doc.add_paragraph("2. Ubah Nama Toko, Alamat, Telepon, Format Nota, Batas Diskon Kasir (%), dan Toggle Stok Minus.")
    doc.add_paragraph("3. Klik 'Simpan Pengaturan Toko'. Pengaturan baru langsung aktif di seluruh sistem.")

    add_screenshot_box("Halaman Pengaturan Toko", "Menampilkan form konfigurasi identitas workshop dan batas diskon POS.")

    # Save
    doc.save("MANUAL_BOOK_RAJAWALI_MOTOR.docx")
    print("MANUAL_BOOK_RAJAWALI_MOTOR.docx created successfully!")

if __name__ == "__main__":
    create_manual_book()
